from docx import Document
from collections import defaultdict
import json
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from utils.generate_doc.flatten_json import flatten_json

def generate_building_registry_docx_simple(json_path: str, ocr_path: str, lang: str) -> Document:
    # JSON 로드
    with open(json_path, 'r', encoding='utf-8') as f:
        replacements = json.load(f)

    doc = Document()  # 문서 객체 생성

    if lang=="일본어":
        serial_number_label = "固有番号"
        date_of_issue_label = "交付日"
        competent_registry_Office_label = "管轄登記所"
        blank_notice_label = "以下余白"
    elif lang=="중국어":
        serial_number_label = "固有番号"
        date_of_issue_label = "发证日期"
        competent_registry_Office_label = "管辖登记机关"
        blank_notice_label = "以下为空白"
    elif lang=="베트남어":
        serial_number_label = "Số định danh"
        date_of_issue_label = "Ngày cấp"
        competent_registry_Office_label = "Cơ quan đăng ký có thẩm quyền"
        blank_notice_label = "Phần còn lại của trang này để trống"
    else:
        serial_number_label = "Serial Number"
        date_of_issue_label = "Date Of Issue"
        competent_registry_Office_label = "Competent Registry Office"
        blank_notice_label = "Nothing follows"
 
    # === 문서 상단 텍스트 (표 위에) 추가 ===
    headings = [
        (f"{replacements['documentType']}", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER),
        (f"- {replacements['typeOfRegistration']} -", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER),
        (f"{serial_number_label} {replacements['serialNumber']}", 11, False, WD_PARAGRAPH_ALIGNMENT.RIGHT),
        (f"[{replacements['typeOfRegistration']}] {replacements['address']}", 11, False, WD_PARAGRAPH_ALIGNMENT.LEFT)
    ]

    for text, size, bold, align in headings:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold

    for t in replacements.get("tables", []):
        header_text = t.get("header", "")
        columns = t.get("columns", [])
        rows_2d = t.get("rows", [])

        # 최소 1열은 보장
        num_cols = max(1, len(columns))
        # 표제(헤더) 1행 + 컬럼헤더 1행 + 본문행 N
        num_rows = 2 + len(rows_2d)

        doc_table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

        # (row 0) 표제부 한 줄 병합 + 텍스트/스타일
        title_cell = doc_table.cell(0, 0)
        if num_cols > 1:
            title_cell = title_cell.merge(doc_table.cell(0, num_cols - 1))
        title_cell.text = header_text
        for para in title_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(12)
                run.bold = True

        # (row 1) 컬럼 헤더
        for c_idx in range(num_cols):
            cell = doc_table.cell(1, c_idx)
            text = str(columns[c_idx]) if c_idx < len(columns) else ""
            cell.text = text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)  # 보통 헤더는 10pt, 굵게 원하면 run.bold=True

        # (row >= 2) 본문 채우기
        for r_idx, row_values in enumerate(rows_2d, start=2):
            # 안전하게 길이 다를 수 있으니 num_cols 기준으로 채움
            for c_idx in range(num_cols):
                cell = doc_table.cell(r_idx, c_idx)
                if c_idx < len(row_values):
                    val = row_values[c_idx]
                    # dict 형태 지원 (e.g., {"text": "..."}), 아니면 문자열/숫자 그대로
                    if isinstance(val, dict):
                        val = val.get("text", "")
                    cell.text = "" if val is None else str(val)
                else:
                    cell.text = ""
                # 폰트 사이즈
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

        # 표 간 간격
        doc.add_paragraph()


    # === 문서 하단 텍스트 (표 아래에) 추가 ===
    p = doc.add_paragraph(f"-- {blank_notice_label} --")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 관할등기소 오른쪽 정렬
    p = doc.add_paragraph(f"{competent_registry_Office_label} {replacements['competentRegistryOffice']}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    remarks_text = '\n'.join(replacements['remarks'])  # 한 문장으로 결합
    p = doc.add_paragraph()
    run = p.add_run(remarks_text)
    run.font.size = Pt(9)

    # 하단
    doc.add_paragraph(f"{date_of_issue_label} : {replacements['dateOfIssue']}")

    return doc