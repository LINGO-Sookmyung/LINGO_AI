from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

# 열 너비 설정 함수
def set_column_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc = row.cells[idx]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(width))  
            tcPr.append(tcW)

# 등록기준지 열 너비 조정
def set_reg_table_widths(table):
    widths = [2500, 7000]  
    for row in table.rows:
        for idx, width in enumerate(widths):
            tc = row.cells[idx]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(width))
            tcPr.append(tcW)

# 가족사항 라벨 열 너비 축소
def set_label_table_width(table, width=2000):
    cell = table.cell(0, 0)._tc
    tcPr = cell.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width))
    tcPr.append(tcW)
def _normalize_replacements(obj):
    """dict 또는 dict 리스트/래핑을 안전하게 dict로 정규화"""
    if isinstance(obj, dict):
        for k in ("data", "payload", "result"):
            if k in obj and isinstance(obj[k], dict):
                return _normalize_replacements(obj[k])
        for k in ("items", "results", "list"):
            if k in obj and isinstance(obj[k], list) and obj[k]:
                first = obj[k][0]
                if isinstance(first, dict):
                    return first
        return obj
    if isinstance(obj, list):
        if not obj:
            raise ValueError("입력 JSON 리스트가 비어 있습니다.")
        if isinstance(obj[0], dict):
            return obj[0]
        raise TypeError("리스트의 첫 요소가 dict가 아닙니다.")
    raise TypeError("입력 JSON은 dict 또는 dict 리스트여야 합니다.")


def generate_family_relationship_docx(json_path: str, lang: str) -> Document:
    # JSON 로드 + 정규화
    with open(json_path, 'r', encoding='utf-8') as f:
        raw = json.load(f)
    replacements = _normalize_replacements(raw)

    doc = Document()

    # 라벨 번역 셋업 (그대로)
    if lang == "일본어":
        original_domicile = "本籍地"
        family_detail = "家族事項"
        time_of_issue_label = "発行日"
        applicant_label = "申請者"
        certificate_number_label = "証明書番号"
    elif lang == "중국어":
        original_domicile = "户籍所在地"
        family_detail = "家庭情况"
        time_of_issue_label = "签发日期"
        applicant_label = "申请人"
        certificate_number_label = "证书编号"
    elif lang == "베트남어":
        original_domicile = "Nơi đăng ký hộ tịch"
        family_detail = "Thông tin gia đình"
        time_of_issue_label = "Ngày cấp"
        applicant_label = "Người nộp đơn"
        certificate_number_label = "Số giấy chứng nhận"
    else:
        original_domicile = "Registered Domicile"
        family_detail = "Family Details"
        time_of_issue_label = "Time of Issue"
        applicant_label = "Applicant"
        certificate_number_label = "Certificate Number"

    # 안전한 getter들
    def gs(key, default=""):
        v = replacements.get(key, default)
        return "" if v is None else str(v)
    def gobj(key):
        v = replacements.get(key, {})
        return v if isinstance(v, dict) else {}
    def glist(key):
        v = replacements.get(key, [])
        return v if isinstance(v, list) else []

    # 제목
    title = doc.add_paragraph(gs("documentType"))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.size = Pt(16)
        title.runs[0].bold = True

    # 등록기준지
    doc.add_paragraph()  # spacer
    reg_table = doc.add_table(rows=1, cols=2)
    reg_table.style = 'Table Grid'
    reg_table.cell(0, 0).text = original_domicile
    reg_table.cell(0, 1).text = gs("placeOfFamilyRegistration")
    for cell in reg_table.row_cells(0):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_reg_table_widths(reg_table)

    # 본인 정보
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'

    # columns 6개로 맞추기
    columns = glist("columns")
    default_cols = ["Category","Full Name","Date of Birth","Reg. No.","Sex","Origin"]
    if not columns: columns = default_cols
    if len(columns) < 6: columns = columns + default_cols[len(columns):]
    if len(columns) > 6: columns = columns[:6]

    hdr = table.rows[0].cells
    for idx, col_name in enumerate(columns):
        if idx < len(hdr):
            hdr[idx].text = str(col_name)

    registrant = gobj("registrant")
    row = table.rows[1].cells
    row[0].text = str(registrant.get("category",""))
    row[1].text = str(registrant.get("fullName",""))
    row[2].text = str(registrant.get("dateOfBirth",""))
    row[3].text = str(registrant.get("residentRegistrationNumber",""))
    row[4].text = str(registrant.get("sex",""))
    row[5].text = str(registrant.get("originOfSurname",""))

    for r in table.rows:
        for cell in r.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_column_widths(table, [1100, 2500, 2500, 3000, 900, 1000])

    # 가족사항 라벨
    doc.add_paragraph()
    label_table = doc.add_table(rows=1, cols=1)
    label_table.style = 'Table Grid'
    label_cell = label_table.cell(0, 0)
    label_cell.text = family_detail
    label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_label_table_width(label_table)

    # 가족 구성
    doc.add_paragraph()
    fam = glist("familyMembers")
    def cat(m): return str(m.get("category",""))
    parents = [m for m in fam if cat(m) in ["Father","Mother","父","母","父亲","母亲","Cha","Mẹ"]]
    spouse  = [m for m in fam if cat(m) in ["Spouse","配偶者","配偶","Người phối ngẫu"]]
    children= [m for m in fam if cat(m) in ["Children","子女","子","Con"]]

    # 부모 표 (헤더 포함)
    fam_table = doc.add_table(rows=max(1, len(parents)) + 1, cols=6)
    fam_table.style = 'Table Grid'
    fam_hdr = fam_table.rows[0].cells
    for i, col_name in enumerate(columns):
        if i < len(fam_hdr):
            fam_hdr[i].text = str(col_name)
    for i, member in enumerate(parents):
        r = fam_table.rows[i + 1].cells
        r[0].text = str(member.get("category",""))
        r[1].text = str(member.get("fullName",""))
        r[2].text = str(member.get("dateOfBirth",""))
        r[3].text = str(member.get("residentRegistrationNumber",""))
        r[4].text = str(member.get("sex",""))
        r[5].text = str(member.get("originOfSurname",""))
    for r in fam_table.rows:
        for cell in r.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_column_widths(fam_table, [1100, 2500, 2500, 3000, 900, 1000])

    # 배우자 표
    if spouse:
        doc.add_paragraph()
        spouse_table = doc.add_table(rows=len(spouse), cols=6)
        spouse_table.style = 'Table Grid'
        for i, member in enumerate(spouse):
            r = spouse_table.rows[i].cells
            r[0].text = str(member.get("category",""))
            r[1].text = str(member.get("fullName",""))
            r[2].text = str(member.get("dateOfBirth",""))
            r[3].text = str(member.get("residentRegistrationNumber",""))
            r[4].text = str(member.get("sex",""))
            r[5].text = str(member.get("originOfSurname",""))
        for r in spouse_table.rows:
            for cell in r.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_column_widths(spouse_table, [1100, 2500, 2500, 3000, 900, 1000])

    # 자녀 표
    if children:
        doc.add_paragraph()
        child_table = doc.add_table(rows=len(children), cols=6)
        child_table.style = 'Table Grid'
        for i, member in enumerate(children):
            r = child_table.rows[i].cells
            r[0].text = str(member.get("category",""))
            r[1].text = str(member.get("fullName",""))
            r[2].text = str(member.get("dateOfBirth",""))
            r[3].text = str(member.get("residentRegistrationNumber",""))
            r[4].text = str(member.get("sex",""))
            r[5].text = str(member.get("originOfSurname",""))
        for r in child_table.rows:
            for cell in r.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_column_widths(child_table, [1100, 2500, 2500, 3000, 900, 1000])

    # 비고/발급일 등
    doc.add_paragraph()
    remarks = glist("remarks")
    if remarks:
        note1 = doc.add_paragraph(str(remarks[0]))
        note1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    issuedDate = doc.add_paragraph()
    run = issuedDate.add_run(gs("dateOfIssue"))
    run.font.size = Pt(12)
    issuedDate.alignment = WD_ALIGN_PARAGRAPH.CENTER

    issuingAuthority = gobj("issuingAuthority")
    org = doc.add_paragraph()
    run = org.add_run(f'{issuingAuthority.get("organization","")} {issuingAuthority.get("authorizedOfficer","")}')
    run.bold = True; run.font.size = Pt(13)
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if len(remarks) > 1:
        note2 = doc.add_paragraph(str(remarks[1]))
        note2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note2.paragraph_format.line_spacing = 1

    doc.add_paragraph()
    issuedTime = doc.add_paragraph(f'{time_of_issue_label} : {gs("timeOfIssue")}')
    applicant = doc.add_paragraph(f'{applicant_label} : {gs("applicant")}')
    issuedTime.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    issuedTime.paragraph_format.space_after = Pt(0)
    applicant.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    applicant.paragraph_format.space_before = Pt(0)

    doc.add_paragraph()
    certificate_number_label = doc.add_paragraph()
    run = certificate_number_label.add_run(f'{certificate_number_label} : {gs("certificateNumber")}')
    run.font.size = Pt(10)
    certificate_number_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    certificate_number_label.paragraph_format.space_after = Pt(0)

    if len(remarks) > 2:
        note3 = doc.add_paragraph()
        run = note3.add_run(str(remarks[2]))
        run.font.size = Pt(10)
        note3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note3.paragraph_format.line_spacing = 1

    return doc