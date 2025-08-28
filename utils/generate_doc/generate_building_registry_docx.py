from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from typing import List, Dict, Any
import json, re

TOP_KEYS = [
    "documentType", "typeOfRegistration", "serialNumber",
    "address", "competentRegistryOffice", "dateOfIssue"
]

def _normalize_structured(obj: Any) -> Dict[str, Any]:
    """dict 또는 dict 리스트를 단일 dict로 정규화."""
    if isinstance(obj, dict):
        for k in ("data", "payload", "result"):
            if isinstance(obj.get(k), dict):
                return _normalize_structured(obj[k])
        for k in ("items", "results", "list", "pages"):
            if isinstance(obj.get(k), list) and obj[k] and isinstance(obj[k][0], dict):
                return _merge_pages(obj[k])
        return dict(obj)
    if isinstance(obj, list) and obj and isinstance(obj[0], dict):
        return _merge_pages(obj)
    return {"tables": [], "remarks": []}

def _merge_pages(pages: List[Dict[str, Any]]) -> Dict[str, Any]:
    out = {k: "" for k in TOP_KEYS}
    out["tables"] = []
    out["remarks"] = []
    for p in pages:
        if not isinstance(p, dict): continue
        for k in TOP_KEYS:
            if not out.get(k) and isinstance(p.get(k), str) and p[k].strip():
                out[k] = p[k].strip()
        if isinstance(p.get("tables"), list):
            out["tables"].extend(p["tables"])
        if isinstance(p.get("remarks"), list):
            for r in p["remarks"]:
                if isinstance(r, str) and r.strip() and r not in out["remarks"]:
                    out["remarks"].append(r)
    return out

def _coerce_legacy_sections(rep: dict) -> dict:
    """partOfTitle / owner 형태를 tables 스키마로 자동 변환."""
    if rep.get("tables"):
        return rep
    tables = []

    pt = rep.get("partOfTitle") or {}
    if isinstance(pt, dict) and (pt.get("columns") or pt.get("rows")):
        header = pt.get("header") or ""
        columns = pt.get("columns") or []
        rows_in = pt.get("rows") or []
        norm_rows = []
        if rows_in and isinstance(rows_in[0], dict):
            order = ["descriptionNo","acceptance","location","buildingDetails","causeOfRegistrationAndOtherInformation"]
            for r in rows_in:
                norm_rows.append([(r.get(k, "") or "") for k in order])
        else:
            norm_rows = rows_in
        tables.append({"header": header, "columns": columns, "rows": norm_rows})

    ow = rep.get("owner") or {}
    if isinstance(ow, dict) and (ow.get("columns") or ow.get("rows")):
        header = ow.get("header") or ""
        columns = ow.get("columns") or []
        rows_in = ow.get("rows") or []
        norm_rows = []
        if rows_in and isinstance(rows_in[0], dict):
            order = ["registeredOwner","registrationNumber","finalShare","ownerAddress","priorityNumber"]
            for r in rows_in:
                norm_rows.append([(r.get(k, "") or "") for k in order])
        else:
            norm_rows = rows_in
        tables.append({"header": header, "columns": columns, "rows": norm_rows})

    if tables:
        rep["tables"] = tables
    return rep

def _normalize_header(h: str) -> str:
    """'표 제 부' → '표제부', 괄호/기호 주변 공백 정리."""
    if not isinstance(h, str):
        return ""
    s = re.sub(r"\s+", " ", h.strip())
    if re.fullmatch(r"[가-힣ㄱ-ㅎㅏ-ㅣ](?:\s[가-힣ㄱ-ㅎㅏ-ㅣ]){1,30}", s):
        s = s.replace(" ", "")
    s = s.replace("【 ", "【").replace(" 】", "】")
    s = s.replace("】 (", "】(").replace("( ", "(").replace(" )", ")")
    return s

def _gs(d: Dict[str, Any], k: str) -> str:
    v = d.get(k, "")
    return "" if v is None else str(v)

def _merge_cont_rows_on_rep(rep: Dict[str, Any]) -> Dict[str, Any]:
    """워드 출력 전 안전망 병합(연속행을 앞행 뒤에 붙이기)."""
    tables = rep.get("tables") or []
    fixed = []
    for t in tables:
        cols = max(len(t.get("columns") or []), max((len(r) for r in (t.get("rows") or [])), default=0))
        rows = [list(r) + [""]*(cols-len(r)) for r in (t.get("rows") or [])]
        out = []
        for r in rows:
            left_empty = all(not (c or "").strip() for c in r[:max(1, cols//2)])
            right_val  = any((c or "").strip() for c in r[max(1, cols//2):])
            if left_empty and right_val and out:
                prev = out[-1]
                for j in range(cols):
                    if (r[j] or "").strip():
                        prev[j] = (prev[j] + ("\n" if prev[j] else "") + r[j]).strip()
            else:
                out.append(r)
        t2 = dict(t); t2["rows"] = out; fixed.append(t2)
    rep["tables"] = fixed
    return rep

def generate_building_registry_docx(json_path: str, ocr_path: str, lang: str) -> Document:
    """
    정책:
    - 표는 오직 rep['tables']만 사용(= 구조화 결과 기반)
    - partOfTitle/owner → tables 변환
    - 헤더를 정규화해서 '표 제 부' 같은 끊김 제거
    - 연속행 병합(워드 출력 전 보정)
    - 하단 고정 블록(빈칸 알림/관할/참고/일시) 삽입
    """
    with open(json_path, encoding="utf-8") as f:
        raw_struct = json.load(f)

    rep = _normalize_structured(raw_struct)
    rep = _coerce_legacy_sections(rep)
    rep = _merge_cont_rows_on_rep(rep)  # 안전망 병합

    # 언어별 라벨
    if lang == "일본어":
        L_SN, L_DOI, L_OFFICE, L_BLANK = "固有番号", "交付日", "管轄登記所", "以下余白"
    elif lang == "중국어":
        L_SN, L_DOI, L_OFFICE, L_BLANK = "固有编号", "发证日期", "管辖登记机关", "以下为空白"
    elif lang == "베트남어":
        L_SN, L_DOI, L_OFFICE, L_BLANK = "Số định danh", "Ngày cấp", "Cơ quan đăng ký có thẩm quyền", "Phần còn lại của trang này để trống"
    else:
        L_SN, L_DOI, L_OFFICE, L_BLANK = "Serial Number", "Date Of Issue", "Competent Registry Office", "Nothing follows"

    doc = Document()

    # 상단 제목
    headings = [
        (_gs(rep, "documentType"), 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER),
        (f"- {_gs(rep, 'typeOfRegistration')} -", 16, True, WD_PARAGRAPH_ALIGNMENT.CENTER),
        (f"{L_SN} {_gs(rep, 'serialNumber')}", 11, False, WD_PARAGRAPH_ALIGNMENT.RIGHT),
        (f"[{_gs(rep, 'typeOfRegistration')}] {_gs(rep, 'address')}", 11, False, WD_PARAGRAPH_ALIGNMENT.LEFT),
    ]
    for text, size, bold, align in headings:
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold

    # 표 렌더링
    rep_tables = rep.get("tables", []) or []
    for rep_tbl in rep_tables:
        header  = _normalize_header(rep_tbl.get("header", ""))
        columns = rep_tbl.get("columns", []) or []
        rows    = rep_tbl.get("rows", []) or []

        cols = max(len(columns), max((len(r) for r in rows), default=0))
        cols = max(cols, 1)

        table = doc.add_table(rows=0, cols=cols, style="Table Grid")

        # 큰 헤더
        if header:
            hdr = table.add_row().cells
            hdr[0].text = header
            for j in range(1, cols):
                hdr[0].merge(hdr[j])
            for para in hdr[0].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(12)
                    run.bold = True

        # 컬럼 헤더
        if columns:
            tr = table.add_row().cells
            for j in range(cols):
                tr[j].text = str(columns[j]) if j < len(columns) else ""
            for c in tr:
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(10)
                    run.bold = True

        # 데이터 행
        for r in rows:
            tr = table.add_row().cells
            for j in range(cols):
                tr[j].text = "" if j >= len(r) or r[j] is None else str(r[j])
            for c in tr:
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(10)

        doc.add_paragraph()  # 표 간 간격

    NOTES_BY_LANG = {
        "일본어": [
            "[ 参考事項 ]",
            "ア. 登記記録で有効な持分を有する所有者または共有者の現況を表示します。",
            "イ. 最終持分は登記名義人が有する最終持分であり、2つ以上の順位番号に持分を有する場合はその持分を合算しました。",
            "ウ. 順位番号は登記名義人を基準として付与された登記の順位番号です。",
            "エ. 申請事項と関係のない所有権（甲区）および所有権以外の権利（乙区事項）は表示していません。",
            "オ. 持分が分割登記された資料は、全体の持分を総合して整理したものです。",
            "＊ 実線で抹消された部分は抹消事項を表示します。＊ 記載事項のない甲区・乙区は『記載事項なし』と表示します。"
        ],
        "중국어": [
            "[ 参考事项 ]",
            "一. 显示在登记记录中拥有有效份额的所有者或共有人现况。",
            "二. 最终份额为登记名义人所拥有的最终份额，如在两个以上的顺序号中拥有份额，则将其合并计算。",
            "三. 顺序号是以登记名义人为基准所赋予的登记顺序号。",
            "四. 与申请事项无关的所有权（甲区）及非所有权的权利（乙区事项）不予显示。",
            "五. 份额被分割登记的资料系将整体份额汇总整理后予以表示。",
            "* 以实线划去的部分表示为注销事项。* 无记载事项的甲区、乙区标注为“无记载事项”。"
        ],
        "베트남어": [
            "[ Ghi chú ]",
            "a. Hiển thị hiện trạng chủ sở hữu hoặc đồng sở hữu có phần sở hữu hợp lệ trong hồ sơ đăng ký.",
            "b. ‘Phần sở hữu cuối cùng’ là phần sở hữu mà người đứng tên đăng ký đang có; nếu có phần sở hữu ở từ 2 số thứ tự trở lên thì được cộng gộp.",
            "c. Số thứ tự được cấp dựa trên người đứng tên đăng ký.",
            "d. Các quyền không liên quan đến nội dung yêu cầu (quyền sở hữu ở Quyển A) và các quyền khác ngoài quyền sở hữu (Quyển B) không được hiển thị.",
            "đ. Trường hợp phần sở hữu được phân chia thì thông tin được tổng hợp theo tổng phần sở hữu.",
            "* Phần kẻ bằng đường liền là nội dung đã bị xóa. * Quyển A/B không có nội dung sẽ được ghi 'Không có nội dung'."
        ],
        # 기본(영어)
        "default": [
            "[ Notes ]",
            "a. Shows the current holders (owners/co-owners) who have valid shares in the registry record.",
            "b. The “final share” is the share ultimately held by the registered owner; if the owner has shares under two or more serial numbers, they are summed.",
            "c. The serial number is the registration sequence number assigned based on the registered owner.",
            "d. Rights not related to the requested matter (ownership in Section A) and rights other than ownership (Section B) are not displayed.",
            "e. If shares have been subdivided, the total share is aggregated and presented.",
            "* Solid line indicates a cancelled item. * Sections A/B with no entries are shown as ‘No entries’.",
        ],
    }

    # 언어별 notes 선택
    if lang in NOTES_BY_LANG:
        notes = NOTES_BY_LANG[lang]
    else:
        notes = NOTES_BY_LANG["default"]

    office_name = _gs(rep, "competentRegistryOffice") 
    date_issued = _gs(rep, "dateOfIssue")              

    p = doc.add_paragraph(f"-- {L_BLANK} --")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    if office_name:
        p = doc.add_paragraph(f"{L_OFFICE} {office_name}")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    notes = NOTES_BY_LANG.get(lang, NOTES_BY_LANG["default"])
    for line in notes:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)

    if date_issued:
        doc.add_paragraph(f"{L_DOI} : {date_issued}")

    return doc