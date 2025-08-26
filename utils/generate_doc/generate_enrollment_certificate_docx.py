from docx import Document
import json
import re

def has_drawing(run):
    return bool(run._element.xpath('.//w:drawing') or run._element.xpath('.//w:pict'))

def replace_in_runs(paragraphs, replacements: dict):
    if not isinstance(replacements, dict):
        raise TypeError("replacements는 dict여야 합니다. (키=플레이스홀더, 값=치환문자열)")

    for para in paragraphs:
        if not para.runs:
            continue
        text_runs, drawing_runs = [], []
        for run in para.runs:
            (drawing_runs if has_drawing(run) else text_runs).append(run)

        # 텍스트 run 병합 후 치환
        buffer = "".join(r.text for r in text_runs)
        for key, val in replacements.items():
            pattern = r"\{\{\s*" + re.escape(str(key)) + r"\s*\}\}"
            buffer = re.sub(pattern, "" if val is None else str(val), buffer)

        # 기존 텍스트 run 제거
        for run in text_runs:
            para._element.remove(run._element)

        # 새 텍스트 run 삽입
        if text_runs:
            first = text_runs[0]
            new_run = para.add_run(buffer)
            new_run.style = first.style
            new_run.font.size = first.font.size
            new_run.bold = first.bold
            new_run.italic = first.italic
            new_run.underline = first.underline

        # 도형 run 원래 순서 유지
        for run in drawing_runs:
            para._element.append(run._element)

def _normalize_replacements(obj):
    """
    - dict 그대로면 통과
    - [dict]이면 첫 요소 사용
    - {"data": dict} / {"items":[dict]} 같은 흔한 래핑도 처리
    """
    if isinstance(obj, dict):
        # 흔한 래핑 해제
        for k in ("data", "payload", "result"):
            if k in obj and isinstance(obj[k], dict):
                return _normalize_replacements(obj[k])
        for k in ("items", "results", "list"):
            if k in obj and isinstance(obj[k], list) and obj[k]:
                first = obj[k][0]
                if isinstance(first, dict):
                    return first
        return obj

    if isinstance(obj, list):
        if not obj:
            raise ValueError("입력 JSON 리스트가 비어 있습니다.")
        if isinstance(obj[0], dict):
            return obj[0]
        raise TypeError("리스트의 첫 요소가 dict가 아닙니다.")

    raise TypeError("입력 JSON은 dict 또는 dict 리스트여야 합니다.")

def generate_enrollment_certificate_docx(json_path: str, lang: str) -> Document:
    # JSON 로드 + 정규화
    with open(json_path, 'r', encoding='utf-8') as f:
        raw = json.load(f)
    replacements = _normalize_replacements(raw)

    # 템플릿 선택
    if lang == "일본어":
        tpl = "templates/japanese_template_enrollment_certificate.docx"
    elif lang == "중국어":
        tpl = "templates/chinese_template_enrollment_certificate.docx"
    elif lang == "베트남어":
        tpl = "templates/vietnamese_template_enrollment_certificate.docx"
    else:
        tpl = "templates/english_template_enrollment_certificate.docx"

    doc = Document(tpl)

    # 본문
    replace_in_runs(doc.paragraphs, replacements)
    #머리말
    for section in doc.sections:
        replace_in_runs(section.header.paragraphs, replacements)

    return doc